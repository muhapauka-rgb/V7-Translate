from datetime import datetime
from html import escape
from pathlib import Path
from typing import Optional

from app.core.errors import UserFacingError
from app.models import AudioFileInfo, ExportOptions, Segment
from app.utils.format_utils import sanitize_filename
from app.utils.time_utils import format_seconds_hhmmss, format_seconds_srt


class ExportService:
    _qt_app = None
    _FONT_FAMILY = "Helvetica Neue"
    _PAGE_MARGIN_CM = 2.0
    _PAGE_MARGIN_MM = 20.0
    _TITLE_SIZE_PT = 18
    _META_SIZE_PT = 10.5
    _BODY_SIZE_PT = 11.5
    _TIMECODE_SIZE_PT = 11
    _TITLE_COLOR = (17, 24, 39)
    _META_COLOR = (71, 85, 105)
    _TIMECODE_COLOR = (100, 116, 139)
    _EN_COLOR = (30, 41, 59)
    _RU_COLOR = (8, 145, 178)

    @staticmethod
    def export_outputs(
        output_dir: Path,
        audio_info: AudioFileInfo,
        segments: list[Segment],
        model_name: str,
        export_options: ExportOptions,
        base_name: Optional[str] = None,
    ) -> list[Path]:
        output_dir.mkdir(parents=True, exist_ok=True)
        safe_base_name = base_name or sanitize_filename(Path(audio_info.file_name).stem)
        generated_files: list[Path] = []

        if export_options.create_docx:
            generated_files.append(
                ExportService.export_docx(
                    output_dir=output_dir,
                    base_name=safe_base_name,
                    audio_info=audio_info,
                    segments=segments,
                    model_name=model_name,
                )
            )
        if export_options.save_pdf:
            generated_files.append(
                ExportService.export_pdf(
                    output_dir=output_dir,
                    base_name=safe_base_name,
                    audio_info=audio_info,
                    segments=segments,
                    model_name=model_name,
                )
            )
        if export_options.save_srt:
            generated_files.append(ExportService.export_srt(output_dir, safe_base_name, segments))
        if export_options.save_txt:
            generated_files.append(ExportService.export_txt(output_dir, safe_base_name, segments))

        return generated_files

    @staticmethod
    def export_txt(output_dir: Path, base_name: str, segments: list[Segment]) -> Path:
        output_path = output_dir / f"{base_name}.txt"

        with output_path.open("w", encoding="utf-8") as file:
            for segment in segments:
                file.write(f"{format_seconds_hhmmss(segment.start)} - {format_seconds_hhmmss(segment.end)}\n")
                file.write(f"EN: {segment.text_en}\n")
                if segment.text_ru:
                    file.write(f"RU: {segment.text_ru}\n")
                file.write("\n")

        return output_path

    @staticmethod
    def export_srt(output_dir: Path, base_name: str, segments: list[Segment]) -> Path:
        output_path = output_dir / f"{base_name}.srt"

        with output_path.open("w", encoding="utf-8") as file:
            for index, segment in enumerate(segments, start=1):
                file.write(f"{index}\n")
                file.write(f"{format_seconds_srt(segment.start)} --> {format_seconds_srt(segment.end)}\n")
                file.write(f"{segment.text_en}\n")
                if segment.text_ru:
                    file.write(f"{segment.text_ru}\n")
                file.write("\n")

        return output_path

    @staticmethod
    def export_docx(
        output_dir: Path,
        base_name: str,
        audio_info: AudioFileInfo,
        segments: list[Segment],
        model_name: str,
    ) -> Path:
        try:
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError as exc:
            raise UserFacingError(
                "Для экспорта в Word не установлен python-docx. Установите зависимости из requirements.txt."
            ) from exc

        output_path = output_dir / f"{base_name}.docx"

        document = Document()
        ExportService._configure_document(document)

        title = document.add_paragraph(style="ALT_Title")
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.add_run(audio_info.file_name)

        for segment in segments:
            document.add_paragraph(
                f"{format_seconds_hhmmss(segment.start)} - {format_seconds_hhmmss(segment.end)}",
                style="ALT_Timecode",
            )
            document.add_paragraph(f"EN: {segment.text_en}", style="ALT_EN")
            if segment.text_ru:
                document.add_paragraph(f"RU: {segment.text_ru}", style="ALT_RU")

        document.save(output_path)
        return output_path

    @staticmethod
    def export_pdf(
        output_dir: Path,
        base_name: str,
        audio_info: AudioFileInfo,
        segments: list[Segment],
        model_name: str,
    ) -> Path:
        try:
            from PySide6.QtCore import QMarginsF, QSizeF
            from PySide6.QtGui import QFont, QPageLayout, QPageSize, QTextDocument
            from PySide6.QtPrintSupport import QPrinter
            from PySide6.QtWidgets import QApplication
        except ImportError as exc:
            raise UserFacingError(
                "Для экспорта в PDF не хватает компонентов PySide6. Переустановите зависимости."
            ) from exc

        if QApplication.instance() is None:
            ExportService._qt_app = QApplication([])

        output_path = output_dir / f"{base_name}.pdf"
        printer = QPrinter(QPrinter.HighResolution)
        printer.setOutputFormat(QPrinter.PdfFormat)
        printer.setOutputFileName(str(output_path))
        printer.setPageSize(QPageSize(QPageSize.A4))
        printer.setFullPage(False)
        printer.setPageMargins(
            QMarginsF(
                ExportService._PAGE_MARGIN_MM,
                ExportService._PAGE_MARGIN_MM,
                ExportService._PAGE_MARGIN_MM,
                ExportService._PAGE_MARGIN_MM,
            ),
            QPageLayout.Millimeter,
        )

        document = QTextDocument()
        default_font = QFont(ExportService._FONT_FAMILY)
        default_font.setPointSizeF(ExportService._BODY_SIZE_PT)
        document.setDefaultFont(default_font)
        document.setDocumentMargin(0)
        page_rect = printer.pageRect(QPrinter.Point)
        document.setPageSize(QSizeF(page_rect.width(), page_rect.height()))
        document.setHtml(ExportService._build_pdf_html(audio_info, segments, model_name))
        document.print_(printer)
        return output_path

    @staticmethod
    def _configure_document(document) -> None:
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_LINE_SPACING
        from docx.shared import RGBColor
        from docx.shared import Cm, Pt

        section = document.sections[0]
        section.top_margin = Cm(ExportService._PAGE_MARGIN_CM)
        section.bottom_margin = Cm(ExportService._PAGE_MARGIN_CM)
        section.left_margin = Cm(ExportService._PAGE_MARGIN_CM)
        section.right_margin = Cm(ExportService._PAGE_MARGIN_CM)

        styles = document.styles
        style_specs = {
            "ALT_Title": {
                "size": ExportService._TITLE_SIZE_PT,
                "bold": True,
                "space_after": 12,
                "line_spacing": 1.0,
                "color": RGBColor(*ExportService._TITLE_COLOR),
            },
            "ALT_Meta": {
                "size": ExportService._META_SIZE_PT,
                "bold": False,
                "space_after": 12,
                "line_spacing": 1.0,
                "color": RGBColor(*ExportService._META_COLOR),
            },
            "ALT_Timecode": {
                "size": ExportService._TIMECODE_SIZE_PT,
                "bold": True,
                "space_after": 6,
                "line_spacing": 1.5,
                "color": RGBColor(*ExportService._TIMECODE_COLOR),
            },
            "ALT_EN": {
                "size": ExportService._BODY_SIZE_PT,
                "bold": False,
                "space_after": 5,
                "line_spacing": 1.5,
                "color": RGBColor(*ExportService._EN_COLOR),
            },
            "ALT_RU": {
                "size": ExportService._BODY_SIZE_PT,
                "bold": False,
                "space_after": 16,
                "line_spacing": 1.5,
                "color": RGBColor(*ExportService._RU_COLOR),
            },
        }

        for style_name, spec in style_specs.items():
            try:
                style = styles[style_name]
            except KeyError:
                style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

            style.font.name = ExportService._FONT_FAMILY
            style.font.size = Pt(spec["size"])
            style.font.bold = spec["bold"]
            style.font.color.rgb = spec["color"]
            style.paragraph_format.space_after = Pt(spec["space_after"])
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            style.paragraph_format.line_spacing = spec["line_spacing"]

    @staticmethod
    def _build_pdf_html(audio_info: AudioFileInfo, segments: list[Segment], model_name: str) -> str:
        title_color = ExportService._rgb_css(ExportService._TITLE_COLOR)
        timecode_color = ExportService._rgb_css(ExportService._TIMECODE_COLOR)
        en_color = ExportService._rgb_css(ExportService._EN_COLOR)
        ru_color = ExportService._rgb_css(ExportService._RU_COLOR)
        blocks = []
        for segment in segments:
            ru_html = ""
            if segment.text_ru:
                ru_html = f"<div class='ru'><b>RU:</b> {escape(segment.text_ru)}</div>"
            blocks.append(
                """
                <div class="segment">
                  <div class="timecode">{start} - {end}</div>
                  <div class="en"><b>EN:</b> {en}</div>
                  {ru}
                </div>
                """.format(
                    start=format_seconds_hhmmss(segment.start),
                    end=format_seconds_hhmmss(segment.end),
                    en=escape(segment.text_en),
                    ru=ru_html,
                )
            )

        return f"""
        <html>
          <head>
            <meta charset="utf-8" />
            <style>
              * {{
                box-sizing: border-box;
              }}
              body {{
                font-family: "{ExportService._FONT_FAMILY}", Arial, sans-serif;
                color: {en_color};
                font-size: {ExportService._BODY_SIZE_PT}pt;
                line-height: 1.5;
                background: #ffffff;
                margin: 0;
                padding: 0;
              }}
              .page {{
                background: #ffffff;
                margin: 0;
                padding: 0;
              }}
              h1 {{
                text-align: center;
                color: {title_color};
                font-size: {ExportService._TITLE_SIZE_PT}pt;
                font-weight: 700;
                margin-top: 0;
                margin-bottom: 12pt;
              }}
              .segment {{
                margin: 0 0 16pt 0;
                padding: 0;
                page-break-inside: avoid;
              }}
              .timecode {{
                color: {timecode_color};
                font-size: {ExportService._TIMECODE_SIZE_PT}pt;
                font-weight: 700;
                line-height: 1.5;
                margin: 0 0 6pt 0;
              }}
              .en {{
                color: {en_color};
                font-size: {ExportService._BODY_SIZE_PT}pt;
                line-height: 1.5;
                margin: 0 0 5pt 0;
              }}
              .ru {{
                color: {ru_color};
                font-size: {ExportService._BODY_SIZE_PT}pt;
                line-height: 1.5;
                margin: 0;
              }}
            </style>
          </head>
          <body>
            <div class="page">
              <h1>{escape(audio_info.file_name)}</h1>
              {''.join(blocks)}
            </div>
          </body>
        </html>
        """

    @staticmethod
    def _rgb_css(rgb: tuple[int, int, int]) -> str:
        return f"rgb({rgb[0]}, {rgb[1]}, {rgb[2]})"
